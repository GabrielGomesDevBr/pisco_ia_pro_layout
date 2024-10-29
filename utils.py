# utils.py
"""
Funções utilitárias do sistema
Contém funções auxiliares usadas em diferentes partes da aplicação
"""

import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime
from constants import USERS_DB

def check_login(username, password):
    """
    Verifica as credenciais de login do usuário
    
    Args:
        username (str): Nome de usuário ou email
        password (str): Senha do usuário
    
    Returns:
        tuple: (bool, str) - (sucesso do login, nome do usuário)
    """
    if username in USERS_DB:
        if USERS_DB[username]['password'] == password:
            return True, USERS_DB[username]['name']
    return False, None

def convert_markdown_to_docx(markdown_text):
    """
    Converte texto markdown para arquivo DOCX
    
    Args:
        markdown_text (str): Texto em formato markdown
    
    Returns:
        io.BytesIO: Arquivo DOCX em memória
    """
    doc = Document()
    
    # Configuração do cabeçalho
    header = doc.add_heading('Relatório Psicológico', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Limpeza e formatação do texto
    cleaned_text = re.sub(r'\*{1,2}', '', markdown_text)  # Remove asteriscos
    cleaned_text = re.sub(r'^#+\s', '', cleaned_text, flags=re.MULTILINE)  # Remove marcadores de título
    cleaned_text = re.sub(r'^\-\s', '• ', cleaned_text, flags=re.MULTILINE)  # Converte hífens em bullets
    
    # Divisão em parágrafos
    paragraphs = cleaned_text.split('\n\n')
    
    # Adiciona cada parágrafo ao documento
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            p.add_run(para.strip())
            p.paragraph_format.space_after = Pt(12)
    
    # Salva o documento em memória
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def create_prompt(report_type, tone, patient_data, specific_fields):
    """
    Cria o prompt para o modelo de IA
    
    Args:
        report_type (str): Tipo de relatório
        tone (str): Tom do relatório
        patient_data (dict): Dados do paciente
        specific_fields (dict): Campos específicos do relatório
    
    Returns:
        str: Prompt formatado
    """
    base_template = f"""
    Você é um assistente especializado em psicologia, focado na geração de {report_type}.
    
    Tom do relatório: {tone}
    
    Dados do paciente:
    Nome: {patient_data['nome']}
    Idade: {patient_data['idade']} anos
    Gênero: {patient_data['genero']}
    Data da avaliação: {patient_data['data_avaliacao']}
    Abordagem terapêutica: {patient_data['abordagem_terapeutica']}

    Informações específicas:
    """
    
    for key, value in specific_fields.items():
        base_template += f"{key}: {value}\n"
    
    base_template += "\nPor favor, gere um relatório profissional e detalhado."
    
    return base_template

def format_date(date):
    """
    Formata uma data para o padrão brasileiro
    
    Args:
        date (datetime): Data a ser formatada
    
    Returns:
        str: Data formatada
    """
    return date.strftime("%d/%m/%Y")

def generate_filename(report_type):
    """
    Gera um nome de arquivo para o relatório
    
    Args:
        report_type (str): Tipo de relatório
    
    Returns:
        str: Nome do arquivo formatado
    """
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"relatorio_{report_type}_{timestamp}.docx"

def validate_patient_data(data):
    """
    Valida os dados do paciente
    
    Args:
        data (dict): Dados do paciente
    
    Returns:
        tuple: (bool, str) - (dados válidos, mensagem de erro)
    """
    required_fields = ['nome', 'idade', 'genero', 'data_avaliacao', 'abordagem_terapeutica']
    
    for field in required_fields:
        if field not in data or not data[field]:
            return False, f"O campo {field} é obrigatório"
            
    if data['idade'] < 0 or data['idade'] > 120:
        return False, "Idade inválida"
        
    return True, ""
